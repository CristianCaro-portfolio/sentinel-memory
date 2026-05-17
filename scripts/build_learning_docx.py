"""Build the pedagogical .docx from the markdown source.

Reads docs/learning/sentinel-memory-aprendizaje.md and produces a clean Word
document. Keeps things simple — no fancy theming, just headings, paragraphs,
lists, code blocks and tables. Run from the repo root:

    python scripts/build_learning_docx.py
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

REPO = Path(__file__).resolve().parent.parent
SRC = REPO / "docs" / "learning" / "sentinel-memory-aprendizaje.md"
OUT = REPO / "docs" / "learning" / "sentinel-memory-aprendizaje.docx"


# --------------------------------------------------------------------------- #
# tiny markdown parser (only the features we use)                             #
# --------------------------------------------------------------------------- #

def _add_inline(p, text: str) -> None:
    """Render **bold**, *italic*, `code`, [text](url) into a python-docx paragraph."""
    pattern = re.compile(
        r"(\*\*[^*]+\*\*)"            # bold
        r"|(\*[^*]+\*)"               # italic
        r"|(`[^`]+`)"                 # inline code
        r"|(\[[^\]]+\]\([^)]+\))"     # link
    )
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith("**"):
            run = p.add_run(token[2:-2]); run.bold = True
        elif token.startswith("`"):
            run = p.add_run(token[1:-1]); run.font.name = "Consolas"
            run.font.color.rgb = RGBColor(0x1F, 0x6F, 0xEB)
        elif token.startswith("["):
            inner = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if inner:
                label = inner.group(1)
                run = p.add_run(label); run.underline = True
                run.font.color.rgb = RGBColor(0x1F, 0x6F, 0xEB)
        else:
            run = p.add_run(token[1:-1]); run.italic = True
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])


def _add_code_block(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def _add_table(doc: Document, header: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h); run.bold = True
    for r, row in enumerate(rows, start=1):
        cells = table.rows[r].cells
        for c, v in enumerate(row):
            cells[c].text = ""
            _add_inline(cells[c].paragraphs[0], v)


def build() -> None:
    text = SRC.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()

    # default paragraph styling: 11pt sans, modest line spacing
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    n = len(lines)
    pending_table: list[list[str]] | None = None
    pending_header: list[str] | None = None

    while i < n:
        raw = lines[i]
        stripped = raw.strip()

        # fenced code block
        if stripped.startswith("```"):
            code_lines: list[str] = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            _add_code_block(doc, code_lines)
            i += 1
            continue

        # markdown table (header + separator + rows)
        if "|" in stripped and i + 1 < n and re.match(r"^\s*\|?\s*[:\-\| ]+$", lines[i + 1]):
            header_cells = [c.strip() for c in stripped.strip("|").split("|")]
            i += 2  # skip separator
            rows: list[list[str]] = []
            while i < n and "|" in lines[i].strip():
                row_cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                # pad short rows
                while len(row_cells) < len(header_cells):
                    row_cells.append("")
                rows.append(row_cells[: len(header_cells)])
                i += 1
            _add_table(doc, header_cells, rows)
            doc.add_paragraph("")  # breathing room
            continue

        # blockquote
        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            run_pre = p.add_run("│ "); run_pre.font.color.rgb = RGBColor(0x58, 0xA6, 0xFF)
            _add_inline(p, stripped[2:])
            i += 1
            continue

        # horizontal rule
        if stripped == "---":
            p = doc.add_paragraph()
            run = p.add_run("─" * 40)
            run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # headings
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=0)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
        elif stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=3)
        elif stripped.startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, stripped[2:])
        elif re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            _add_inline(p, re.sub(r"^\d+\.\s", "", stripped))
        elif stripped == "":
            doc.add_paragraph("")
        else:
            p = doc.add_paragraph()
            _add_inline(p, stripped)

        i += 1

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"[build_learning_docx] wrote {OUT.relative_to(REPO)}")


if __name__ == "__main__":
    build()
